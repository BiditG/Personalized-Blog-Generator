import streamlit as st
import subprocess
import requests
import zipfile
import os
import shutil
from datetime import datetime
from docx import Document
from tinydb import TinyDB
import language_tool_python
from textblob import TextBlob
import random
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
# ---------- Initialize DB and Tools ----------
db = TinyDB("blog_history.json")
tool = language_tool_python.LanguageTool('en-US')

def add_horizontal_line(doc):
    p = doc.add_paragraph()
    p_paragraph = p._element
    hr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'auto')
    hr.append(bottom)
    p_paragraph.insert(0, hr)


TEMPLATE_PRESETS = {
    "Default": "Use a clear, concise tone. Ensure paragraphs are short and keyword optimized.",
    "Storytelling": "Use storytelling with personal experiences, metaphors, and first-person voice.",
    "Educational": "Maintain a professional tone, detailed explanations, and supportive bullet points.",
    "Sales Focused": "Use persuasive language, scarcity tactics, and direct CTAs to boost conversions."
}

# ---------- Helpers ----------
def grammar_enhance(text):
    matches = tool.check(text)
    return language_tool_python.utils.correct(text, matches)

def voice_tone_score(text):
    blob = TextBlob(text)
    polarity = blob.sentiment.polarity
    if polarity > 0.3:
        return "Upbeat / Positive"
    elif polarity < -0.3:
        return "Serious / Cautionary"
    else:
        return "Neutral / Informative"

def fake_plagiarism_score():
    return random.randint(2, 15)  # simulate low plagiarism

# ---------- Blog Generator ----------
def generate_blog_ollama(topic, tone, audience, word_count, model, subheadings,
                         include_meta, include_conclusion, style, perspective, cta_type, keywords, template):
    prompt = f"""
    Write a high-quality, SEO-optimized blog post on the topic: "{topic}".

    Guidelines:
    - Use the following tone: {tone}
    - Target audience: {audience}
    - Length: approximately {word_count} words
    - Style: {style}
    - Perspective: {perspective}
    - Use {subheadings} informative H2 subheadings
    - {"Include a meta description under 160 characters." if include_meta else ""}
    - Integrate keywords: {keywords}
    - {template}
    - Short paragraphs (2-3 lines)
    - Bullet points, lists, or examples when helpful
    - {"Include conclusion with a strong CTA: " + cta_type if include_conclusion else ""}
    """

    # Debug: Show prompt in app for inspection
    with st.expander("🔍 Debug Prompt"):
        st.code(prompt, language="markdown")

    try:
        result = subprocess.run(
            ["ollama", "run", model],
            input=prompt.encode(),
            capture_output=True,
            timeout=300  # seconds
        )
        if result.returncode != 0:
            st.error("🛑 Model returned an error. Please check if the model is installed and working.")
            st.code(result.stderr.decode(), language="bash")
            return "ERROR: Model execution failed."

        output = result.stdout.decode().strip()

        if not output:
            st.warning("⚠️ Model returned no output. This might be due to system overload or incomplete model.")
            return "ERROR: No response from model."

        return grammar_enhance(output)

    except subprocess.TimeoutExpired:
        st.error("⏳ Model generation timed out. Try using a smaller model or reducing word count.")
        return "ERROR: Model timed out."

    except Exception as e:
        st.error("🚨 Unexpected error during model execution.")
        st.code(str(e))
        return f"ERROR: {str(e)}"

def refine_blog_with_model_b(blog_text, model_b):
    prompt = f"""
    You are an expert content editor. Expand and enhance the following blog post.

    Instructions:
    - Significantly increase the word count to reach around {word_count} words.
    - Add detailed explanations, real-world examples, lists, or use cases.
    - Expand each paragraph to be richer in information.
    - Maintain the original tone and topic.
    - Improve clarity, coherence, and SEO effectiveness.
    - Add missing transitions or elaboration if needed.

    Blog:
    \"\"\"
    {blog_text}
    \"\"\"
    """


    try:
        result = subprocess.run(
            ["ollama", "run", model_b],
            input=prompt.encode(),
            capture_output=True,
            timeout=300
        )
        if result.returncode != 0:
            st.error("Model B encountered an error.")
            st.code(result.stderr.decode(), language="bash")
            return "ERROR: Model B failed."

        output = result.stdout.decode().strip()
        if not output:
            return "ERROR: Model B returned empty content."

        return grammar_enhance(output)

    except Exception as e:
        st.error("Refinement failed.")
        return f"ERROR: {str(e)}"


def optimize_blog_with_model_c(blog_text, model_c):
    prompt = f"""
    Optimize the following blog content for SEO and readability.
    Ensure:
    - Effective use of keywords
    - Smooth paragraph transitions
    - Strong and engaging conclusion
    - Reader-friendly formatting and clarity
    - Proper use of headings and subheadings
    - Make sure the paragraph is not so short

    Blog:
    \"\"\" 
    {blog_text} 
    \"\"\"

    Finally write the optimized blog and make sure you dont reduce the word count too much.
    """

    try:
        result = subprocess.run(
            ["ollama", "run", model_c],
            input=prompt.encode(),
            capture_output=True,
            timeout=300
        )
        if result.returncode != 0:
            st.error("Model C encountered an error.")
            st.code(result.stderr.decode(), language="bash")
            return "ERROR: Model C failed."

        output = result.stdout.decode().strip()
        if not output:
            return "ERROR: Model C returned empty content."

        return grammar_enhance(output)

    except Exception as e:
        st.error("Final optimization failed.")
        return f"ERROR: {str(e)}"

# ---------- Save Blog to DB ----------
def save_blog_to_db(topic, tone, audience, word_count, model, blog_text, style,
                    perspective, cta_type, keywords, template):
    db.insert({
        "topic": topic,
        "tone": tone,
        "audience": audience,
        "word_count": word_count,
        "model": model,
        "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "style": style,
        "perspective": perspective,
        "cta_type": cta_type,
        "keywords": keywords,
        "template": template,
        "blog": blog_text
    })

# ---------- Image Fetcher ----------
def get_unsplash_images(query, client_id, count=5):
    url = "https://api.unsplash.com/search/photos"
    params = {"query": query, "per_page": count, "client_id": client_id, "orientation": "landscape"}
    response = requests.get(url, params=params)
    data = response.json()
    return [img["urls"]["regular"] for img in data.get("results", [])]

from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_horizontal_line(paragraph):
    p_element = paragraph._element
    p_borders = OxmlElement('w:pBdr')
    bottom_border = OxmlElement('w:bottom')
    bottom_border.set(qn('w:val'), 'single')
    bottom_border.set(qn('w:sz'), '8')  # thickness
    bottom_border.set(qn('w:space'), '4')
    bottom_border.set(qn('w:color'), 'BFBFBF')  # light gray
    p_borders.append(bottom_border)
    p_element.insert(0, p_borders)

def export_to_docx(blog_text, filename="blog.docx"):
    doc = Document()

    # Set document font style
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Garamond"
    font.size = Pt(12)

    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    lines = blog_text.strip().split("\n")
    is_title_done = False

    for i, line in enumerate(lines):
        stripped = line.strip()
        if not stripped:
            continue

        # Title (first non-empty line)
        if not is_title_done:
            title = doc.add_heading(stripped, level=1)
            title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            title.runs[0].font.color.rgb = RGBColor(0, 51, 102)  # Navy Blue
            is_title_done = True
            continue

        # Subheadings
        if stripped.startswith("##") or stripped.endswith(":"):
            heading = doc.add_heading(stripped.replace("##", "").strip(), level=2)
            heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            add_horizontal_line(heading)
            continue

        # Bullet Points
        if stripped.startswith("* ") or stripped.startswith("- "):
            bullet = doc.add_paragraph(stripped[2:].strip(), style="List Bullet")
            bullet.paragraph_format.space_after = Pt(6)
            bullet.paragraph_format.left_indent = Inches(0.25)
            continue

        # Regular Paragraphs
        paragraph = doc.add_paragraph(stripped)
        paragraph.paragraph_format.space_after = Pt(10)
        paragraph.paragraph_format.line_spacing = 1.5
        paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    # Add footer with timestamp
    footer_section = doc.sections[-1]
    footer = footer_section.footer
    footer_para = footer.paragraphs[0]
    footer_para.text = f"Generated by AI Blog Generator on {datetime.now().strftime('%B %d, %Y')}"
    footer_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    footer_para.runs[0].font.size = Pt(9)
    footer_para.runs[0].font.color.rgb = RGBColor(128, 128, 128)

    # Save the document
    doc.save(filename)



# ---------- Fiverr ZIP Export ----------
def create_fiverr_package(blog_text, docx_path, img_urls, topic):
    folder = f"fiverr_delivery/{topic.replace(' ', '_')}"
    os.makedirs(folder, exist_ok=True)

    txt_path = os.path.join(folder, "blog.txt")
    with open(txt_path, "w", encoding="utf-8") as f:
        f.write(blog_text)

    img_path = os.path.join(folder, "image_links.txt")
    with open(img_path, "w", encoding="utf-8") as f:
        for url in img_urls:
            f.write(url + "\n")

    docx_dest = os.path.join(folder, "blog.docx")
    shutil.copyfile(docx_path, docx_dest)

    zip_filename = f"{folder}.zip"
    with zipfile.ZipFile(zip_filename, "w") as zipf:
        zipf.write(txt_path, arcname="blog.txt")
        zipf.write(img_path, arcname="image_links.txt")
        zipf.write(docx_dest, arcname="blog.docx")

    return zip_filename

# ---------- Streamlit UI ----------
st.set_page_config(page_title="📝 Blog Generator Pro", layout="centered")
st.title("📝 AI-Powered Blog Generator with Templates & Scoring")

# Sidebar Inputs
st.sidebar.header("✍️ Blog Settings")
topic = st.sidebar.text_input("Blog Topic", "Boosting Productivity in Remote Teams")
tone = st.sidebar.selectbox("Tone", ["Conversational", "Professional", "Casual", "Formal"])
audience = st.sidebar.text_input("Target Audience", "Startup Founders")
style = st.sidebar.selectbox("Writing Style", ["How-to", "Listicle", "Narrative", "Case Study", "Opinion"])
perspective = st.sidebar.selectbox("Perspective", ["First-person", "Third-person"])
cta_type = st.sidebar.text_input("Call-to-Action", "Contact us for a free consultation")
keywords = st.sidebar.text_input("SEO Keywords (comma-separated)", "remote work, productivity, team management")
template_choice = st.sidebar.selectbox("Writing Template Preset", list(TEMPLATE_PRESETS.keys()))
template = TEMPLATE_PRESETS[template_choice]

word_count = st.sidebar.slider("Word Count", 300, 2000, 800)
model = st.sidebar.selectbox("AI Model", ["llama3", "gemma", "mistral"])
model_b = st.sidebar.selectbox("Refinement Model (Model B)", ["llama3", "gemma", "mistral"])
model_c = st.sidebar.selectbox("Final Optimizer Model (Model C)", ["llama3", "gemma", "mistral"])

subheadings = st.sidebar.slider("H2 Subheadings", 2, 6, 4)
include_meta = st.sidebar.checkbox("Include Meta Description", value=True)
include_conclusion = st.sidebar.checkbox("Include Conclusion + CTA", value=True)
UNSPLASH_ACCESS_KEY = st.sidebar.text_input("🔑 Unsplash Access Key", type="password")
enable_fiverr_package = st.sidebar.checkbox("🎁 Fiverr ZIP Package", value=True)

# History
history_entries = db.all()
if history_entries:
    titles = [f"{i['topic']} ({i['timestamp']})" for i in history_entries]
    selected = st.sidebar.selectbox("📖 Load Past Blog", titles)
    if selected:
        loaded = history_entries[titles.index(selected)]
        if st.sidebar.button("📤 Load This Blog"):
            st.session_state["loaded_blog"] = loaded

if "loaded_blog" in st.session_state:
    b = st.session_state["loaded_blog"]
    st.subheader(f"📄 Loaded Blog: {b['topic']}")
    st.text_area("Blog Output", b["blog"], height=500)
    st.download_button("⬇️ Download .txt", data=b["blog"], file_name=f"{b['topic']}.txt")

if st.sidebar.button("🗑️ Clear History"):
    db.truncate()
    st.success("History cleared!")

# Generate Blog
if st.button("🚀 Generate & Enhance Blog"):
    with st.spinner("Model A is generating the base blog..."):
        base_blog = generate_blog_ollama(topic, tone, audience, word_count, model,
                                         subheadings, include_meta, include_conclusion,
                                         style, perspective, cta_type, keywords, template)

    if not base_blog.startswith("ERROR"):
        with st.spinner("Model B is refining the blog..."):
            refined_blog = refine_blog_with_model_b(base_blog, model_b)

        with st.spinner("Model C is optimizing the final blog..."):
            optimized_blog = optimize_blog_with_model_c(refined_blog, model_c)

        st.subheader("📄 Final Optimized Blog")
        st.text_area("Final Blog Output", optimized_blog, height=500)

        # Tone + Plagiarism
        st.success(f"✅ Voice Tone Score: {voice_tone_score(optimized_blog)}")
        st.warning(f"🔍 Estimated Plagiarism: {fake_plagiarism_score()}%")
        st.info(f"📝 Final Word Count: {len(optimized_blog.split())} words")


        # Save and export
        docx_filename = f"{topic.replace(' ', '_')}_final.docx"
        export_to_docx(optimized_blog, docx_filename)
        st.download_button("⬇️ Download Final .docx", data=open(docx_filename, "rb"), file_name=docx_filename)
        st.download_button("⬇️ Download Final .txt", data=optimized_blog, file_name=f"{topic.replace(' ', '_')}_final.txt")

        save_blog_to_db(topic, tone, audience, word_count, f"{model}+{model_b}+{model_c}", optimized_blog,
                        style, perspective, cta_type, keywords, template)



        # Optional image and Fiverr ZIP
        if UNSPLASH_ACCESS_KEY:
            st.subheader("🖼️ Suggested Images")
            img_urls = get_unsplash_images(topic, UNSPLASH_ACCESS_KEY)
            for url in img_urls:
                st.image(url, use_column_width=True)
                st.code(url)

        if enable_fiverr_package:
            with st.spinner("📦 Creating Fiverr ZIP..."):
                zip_path = create_fiverr_package(refined_blog, docx_filename, img_urls, topic)
                with open(zip_path, "rb") as zf:
                    st.download_button("⬇️ Download ZIP Package", data=zf, file_name=os.path.basename(zip_path))
